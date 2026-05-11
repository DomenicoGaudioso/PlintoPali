# -*- coding: utf-8 -*-
import io
from docx import Document
from docx.shared import Inches, RGBColor, Pt
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from datetime import datetime
from src import figura_mesh_fem, genera_verifiche_df # Import new functions


def plot_comparativo_reazioni_mpl(risultati, tipo_confronto='rigido_vs_fem'):
    """
    Genera un grafico a barre comparativo con Matplotlib.
    Restituisce un buffer di byte contenente l'immagine PNG.
    """
    if tipo_confronto == 'rigido_vs_fem':
        dati1 = risultati['statico']
        dati2 = risultati['statico_fem']
        label1 = 'Plinto Rigido'
        label2 = 'Plinto Flessibile (FEM)'
        titolo = 'Confronto Reazioni Statiche: Rigido vs Flessibile'
    else:  # statico_vs_sismico
        dati1 = risultati['statico']
        dati2 = risultati['sismico']
        label1 = 'Statico'
        label2 = 'Sismico'
        titolo = 'Confronto Reazioni: Statico vs Sismico (Rigido)'

    n_pali = len(dati1['x'])
    labels_pali = [f"P{i+1}" for i in range(n_pali)]
    valori1 = np.round(dati1['R'], 1)
    valori2 = np.round(dati2['R'], 1)

    x = np.arange(len(labels_pali))
    width = 0.35

    fig, ax = plt.subplots(figsize=(12, 7))
    rects1 = ax.bar(x - width/2, valori1, width, label=label1, color='steelblue')
    rects2 = ax.bar(x + width/2, valori2, width, label=label2, color='darkorange')

    ax.set_ylabel('Reazione Assiale [kN]')
    ax.set_title(titolo)
    ax.set_xticks(x)
    ax.set_xticklabels(labels_pali)
    ax.legend()
    ax.grid(axis='y', linestyle='--', alpha=0.7)

    ax.bar_label(rects1, padding=3, fmt='%.0f')
    ax.bar_label(rects2, padding=3, fmt='%.0f')

    qamm = risultati.get('Qamm_effettiva_palo', 0)
    if qamm > 0:
        ax.axhline(y=qamm, color='red', linestyle='--', linewidth=2,
                   label=f'Qamm (Gruppo) = {qamm:.0f} kN')
        ax.legend()

    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def add_df_to_doc(doc, df: pd.DataFrame, font_size: int = 9):
    """Aggiunge un DataFrame pandas a un documento docx."""
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    t.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        t.cell(0, j).text = col_name
    for i, row in enumerate(df.itertuples(), start=1):
        for j, val in enumerate(row[1:]):
            cell = t.cell(i, j)
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def plot_stratigrafia_mpl(risultati):
    """
    Genera un grafico della stratigrafia con Matplotlib.
    Restituisce un buffer di byte contenente l'immagine PNG.
    """
    df = risultati['stratigrafia'].copy()

    fig, ax = plt.subplots(figsize=(5, 7))
    colors = ['#D2B48C', '#DEB887', '#F4A460', '#CD853F', '#A0522D', '#8B4513']

    for i, row in df.iterrows():
        z_t, z_b = -row['z_top_m'], -row['z_bot_m']
        c = colors[i % len(colors)]

        ax.fill_betweenx([z_t, z_b], 0, 1, color=c, edgecolor='black', linewidth=0.5)

        testo = (f"Sp: {row['spessore_m']:.1f}m\n"
                 f"φ={row['phi_deg']:.0f}°, cᵤ={row['cu_kPa']:.0f} kPa\n"
                 f"E_ed={row['E_ed_kPa']:.0f} kPa")
        ax.text(0.5, (z_t + z_b) / 2, testo,
                ha='center', va='center', color='black', fontsize=8)

    ax.set_xlim(0, 1)
    ax.set_xticks([])
    ax.set_ylabel("Profondità Z [m]")
    ax.set_title("Profilo Stratigrafico")
    ax.grid(False)

    fig.tight_layout(pad=0.5)

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def crea_report_word(dati, risultati, tabella_pali_df, verifiche_df: pd.DataFrame):
    """
    Crea un documento Word con i risultati dell'analisi.
    """
    doc = Document()

    # Imposta uno stile sobrio con testo nero per i titoli e i paragrafi.
    styles = doc.styles
    black = RGBColor(0, 0, 0)

    # Stile per il testo normale
    style_normal = styles['Normal']
    style_normal.font.color.rgb = black

    # Stile per i titoli di capitolo e il titolo principale
    style_names = ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']
    for style_name in style_names:
        try:
            styles[style_name].font.color.rgb = black
        except KeyError:
            continue  # Stile non presente nel template di base, si procede oltre

    doc.add_heading('Relazione di Calcolo - Fondazione su Pali', 0) # Level 0 is Title
    doc.add_paragraph(f"Software: PlintoPali Ultimate")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Normativa di Riferimento: NTC 2018, Eurocodice 7")
    doc.add_paragraph(f"Disclaimer: Il presente documento è un output di calcolo e non sostituisce la progettazione e la supervisione di un tecnico abilitato.")

    doc.add_heading('1. Dati di Input', level=1) # Level 1 is Heading 1
    doc.add_paragraph(f"Geometria Plinto: {dati.B:.2f}m x {dati.L:.2f}m, H={dati.spessore_plinto:.2f}m")
    doc.add_paragraph(f"Modulo E Cls: {dati.E_cls_MPa:.0f} MPa")
    doc.add_paragraph(f"Disposizione Pali: {dati.n_x} x {dati.n_y}, interasse X={dati.interasse_x:.2f}m, interasse Y={dati.interasse_y:.2f}m")
    doc.add_paragraph(f"Palo: D={dati.diametro_palo:.2f}m, L={dati.lunghezza_palo:.2f}m")
    doc.add_paragraph(f"Azioni: N={dati.N:.0f} kN, Mx={dati.Mx:.0f} kNm, My={dati.My:.0f} kNm")
    doc.add_paragraph(f"Parametri Geotecnici: Nq={dati.Nq:.1f}, Nc={dati.Nc:.1f}, β={dati.beta:.2f}, α={dati.alpha:.2f}, FS={dati.gamma_sicurezza:.2f}")
    doc.add_paragraph(f"Profondità Falda: {dati.falda:.2f} m")

    doc.add_heading('2. Dati e Risultati Geotecnici', level=1)
    doc.add_paragraph("La stratigrafia di input è la seguente:")
    strat_df = risultati['stratigrafia'].copy()
    strat_df_report = strat_df[['spessore_m', 'gamma_dry', 'gamma_sat', 'phi_deg', 'cu_kPa', 'E_ed_kPa']].rename(columns={
        'spessore_m': 'Spessore [m]',
        'gamma_dry': 'γ_dry [kN/m³]',
        'gamma_sat': 'γ_sat [kN/m³]',
        'phi_deg': 'φ [°]',
        'cu_kPa': 'cᵤ [kPa]',
        'E_ed_kPa': 'E_ed [kPa]'
    })
    add_df_to_doc(doc, strat_df_report)

    img_buf_strat = plot_stratigrafia_mpl(risultati)
    doc.add_picture(img_buf_strat, width=Inches(3.5))

    doc.add_paragraph("\nSintesi dei risultati geotecnici:")
    doc.add_paragraph(f"Portata Ultima Palo Singolo (Qult): {risultati['Qult_palo']:.0f} kN")
    doc.add_paragraph(f"Portata Ammissibile Palo Singolo (Qamm): {risultati['Qamm_palo']:.0f} kN")
    eff = risultati['efficienza_gruppo']
    doc.add_paragraph(f"Efficienza Gruppo (η): {eff['eta']:.2f} ({eff['stato']})")
    doc.add_paragraph(f"Portata Ammissibile di Gruppo (Qamm,eff): {risultati['Qamm_effettiva_palo']:.0f} kN")

    doc.add_heading('3. Risultati Analisi', level=1)
    doc.add_paragraph("La tabella seguente riassume le reazioni sui pali calcolate con i diversi modelli.")
    add_df_to_doc(doc, tabella_pali_df)

    doc.add_heading('4. Grafici Comparativi', level=1)

    doc.add_paragraph("Confronto tra modello a plinto rigido e flessibile (FEM) in condizioni statiche.")
    img_buf_1 = plot_comparativo_reazioni_mpl(risultati, 'rigido_vs_fem')
    doc.add_picture(img_buf_1, width=Inches(6.5))

    # Si potrebbero aggiungere altri grafici qui (es. statico vs sismico)

    # Salva il documento in un buffer di memoria
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()