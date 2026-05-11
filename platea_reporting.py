# -*- coding: utf-8 -*-
import io
from datetime import datetime

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt


def add_df_to_doc(doc, df: pd.DataFrame, font_size: int = 9):
    """Aggiunge un DataFrame pandas a un documento docx."""
    df = df.round(2)
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    t.style = 'Table Grid'
    for j, col_name in enumerate(df.columns):
        t.cell(0, j).text = col_name
    for i, row in enumerate(df.itertuples(), start=1):
        for j, val in enumerate(row[1:]):
            cell = t.cell(i, j)
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def plot_risultati_platea_mpl(risultati, z_key, title):
    """Genera un grafico a contorni con Matplotlib."""
    fig, ax = plt.subplots(figsize=(8, 7))

    is_moment = 'M' in z_key
    x = risultati['x_coords']
    y = risultati['y_coords']
    z = risultati[z_key]

    if is_moment:
        x = (risultati['x_coords'][:-1] + risultati['x_coords'][1:]) / 2
        y = (risultati['y_coords'][:-1] + risultati['y_coords'][1:]) / 2

    # Matplotlib's contourf expects Z with shape (Y.shape[0], X.shape[0])
    # I dati da OpenSees sono (ny, nx), che corrisponde a (len(y), len(x))
    # quindi non serve la trasposizione se le coordinate sono corrette.
    contour = ax.contourf(x, y, z, cmap='viridis', levels=15)
    cbar = fig.colorbar(contour, ax=ax)
    
    try:
        cbar.set_label(title.split('[')[1].split(']')[0])
    except IndexError:
        cbar.set_label(title)

    ax.contour(x, y, z, colors='k', linewidths=0.5, levels=15)

    ax.set_title(title)
    ax.set_xlabel('X [m]')
    ax.set_ylabel('Y [m]')
    ax.set_aspect('equal', adjustable='box')
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def crea_report_word_platea(dati_stat, dati_sis, risultati_stat, risultati_sis, kh: float, kv: float):
    """Crea un documento Word con i risultati dell'analisi della platea."""
    doc = Document()
    doc.add_heading('Relazione di Calcolo - Platea di Fondazione (FEM)', 0)
    doc.add_paragraph(f"Software: PlateaFEM")
    doc.add_paragraph(f"Data: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Dati di Input', level=1)
    doc.add_paragraph(f"Dimensioni platea: {dati_stat.B:.2f}m (X) x {dati_stat.L:.2f}m (Y)")
    doc.add_paragraph(f"Spessore: {dati_stat.spessore:.2f} m")
    doc.add_paragraph(f"Modulo E calcestruzzo: {dati_stat.E_cls_MPa:.0f} MPa")
    doc.add_paragraph(f"Modulo di Winkler: {dati_stat.k_winkler_kPa_m:.0f} kPa/m")
    doc.add_paragraph(f"Dimensione mesh: {dati_stat.mesh_size:.2f} m")
    if dati_stat.q_distribuito_kPa != 0:
        doc.add_paragraph(f"Carico distribuito (statico): {dati_stat.q_distribuito_kPa:.2f} kPa")
    
    doc.add_heading('Coefficienti Sismici', level=2)
    doc.add_paragraph(f"kh = {kh:.3f}")
    doc.add_paragraph(f"kv = {kv:.3f}")

    if not dati_stat.pilastri_df.empty:
        doc.add_heading('Carichi sui Pilastri (Statici)', level=2)
        add_df_to_doc(doc, dati_stat.pilastri_df)

    doc.add_heading('2. Risultati di Sintesi', level=1)
    
    # Sintesi Statica
    doc.add_heading('Condizione Statica', level=2)
    cedimento_max_stat = risultati_stat['cedimenti_mm'].max()
    pressione_max_stat = risultati_stat['pressioni_kPa'].max()
    mxx_max_stat = risultati_stat['Mxx_kNm_m'].max()
    myy_max_stat = risultati_stat['Myy_kNm_m'].max()
    doc.add_paragraph(f"Cedimento massimo: {cedimento_max_stat:.2f} mm")
    doc.add_paragraph(f"Pressione massima sul terreno: {pressione_max_stat:.1f} kPa")
    doc.add_paragraph(f"Momento Mxx massimo: {mxx_max_stat:.1f} kNm/m")
    doc.add_paragraph(f"Momento Myy massimo: {myy_max_stat:.1f} kNm/m")

    # Sintesi Sismica
    doc.add_heading('Condizione Sismica', level=2)
    cedimento_max_sis = risultati_sis['cedimenti_mm'].max()
    pressione_max_sis = risultati_sis['pressioni_kPa'].max()
    mxx_max_sis = risultati_sis['Mxx_kNm_m'].max()
    myy_max_sis = risultati_sis['Myy_kNm_m'].max()
    doc.add_paragraph(f"Cedimento massimo: {cedimento_max_sis:.2f} mm")
    doc.add_paragraph(f"Pressione massima sul terreno: {pressione_max_sis:.1f} kPa")
    doc.add_paragraph(f"Momento Mxx massimo: {mxx_max_sis:.1f} kNm/m")
    doc.add_paragraph(f"Momento Myy massimo: {myy_max_sis:.1f} kNm/m")

    doc.add_heading('3. Mappe dei Risultati', level=1)

    # Grafici Statici
    doc.add_heading('Condizione Statica', level=2)
    doc.add_paragraph("Mappa dei cedimenti:")
    img_ced_stat = plot_risultati_platea_mpl(risultati_stat, 'cedimenti_mm', 'Cedimenti [mm] - Statico')
    doc.add_picture(img_ced_stat, width=Inches(6.0))
    doc.add_paragraph("Mappa delle pressioni sul terreno:")
    img_press_stat = plot_risultati_platea_mpl(risultati_stat, 'pressioni_kPa', 'Pressioni [kPa] - Statico')
    doc.add_picture(img_press_stat, width=Inches(6.0))
    doc.add_paragraph("Mappa dei momenti flettenti Mxx:")
    img_mxx_stat = plot_risultati_platea_mpl(risultati_stat, 'Mxx_kNm_m', 'Momento Mxx [kNm/m] - Statico')
    doc.add_picture(img_mxx_stat, width=Inches(6.0))

    # Grafici Sismici
    doc.add_heading('Condizione Sismica', level=2)
    doc.add_paragraph("Mappa dei cedimenti:")
    img_ced_sis = plot_risultati_platea_mpl(risultati_sis, 'cedimenti_mm', 'Cedimenti [mm] - Sismico')
    doc.add_picture(img_ced_sis, width=Inches(6.0))
    doc.add_paragraph("Mappa delle pressioni sul terreno:")
    img_press_sis = plot_risultati_platea_mpl(risultati_sis, 'pressioni_kPa', 'Pressioni [kPa] - Sismico')
    doc.add_picture(img_press_sis, width=Inches(6.0))
    doc.add_paragraph("Mappa dei momenti flettenti Mxx:")
    img_mxx_sis = plot_risultati_platea_mpl(risultati_sis, 'Mxx_kNm_m', 'Momento Mxx [kNm/m] - Sismico')
    doc.add_picture(img_mxx_sis, width=Inches(6.0))

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()