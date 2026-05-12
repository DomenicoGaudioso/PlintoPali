# -*- coding: utf-8 -*-
"""
report.py — Generazione Relazione Tecnica Plinto su Pali
Formati: Word (.docx)
"""
from __future__ import annotations

import io
import math
import tempfile
from datetime import date
from pathlib import Path
from typing import Dict, Any

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src import genera_verifiche_df, tabella_pali_comparativa, tabella_sintesi


def _fig_to_png_bytes(fig) -> bytes:
    """Converte una figura Matplotlib in bytes PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def _plot_geometria_matplotlib(d: Any, r: Dict[str, Any]) -> bytes:
    """Genera un plot statico della geometria del plinto e dei pali con reazioni."""
    rr = r['statico']
    R_vals, x_pali, y_pali = rr['R'], rr['x'], rr['y']

    fig, ax = plt.subplots(figsize=(8, 6))

    # Plinto
    xs_plinto = [-d.B / 2, d.B / 2, d.B / 2, -d.B / 2, -d.B / 2]
    ys_plinto = [-d.L / 2, -d.L / 2, d.L / 2, d.L / 2, -d.L / 2]
    ax.plot(xs_plinto, ys_plinto, color='steelblue', linewidth=2, label='Plinto')
    ax.fill(xs_plinto, ys_plinto, color='lightblue', alpha=0.3)

    # Pali
    R_max_abs = np.max(np.abs(R_vals))
    if R_max_abs < 1e-9: R_max_abs = 1.0 # Evita divisione per zero

    for i in range(len(x_pali)):
        # Disegna il palo come un cerchio
        circle = plt.Circle((x_pali[i], y_pali[i]), d.diametro_palo / 2,
                            facecolor='gray', alpha=0.7, edgecolor='black', linewidth=0.8)
        ax.add_patch(circle)

        # Colore basato sulla reazione (rosso per compressione, blu per trazione)
        norm_R = R_vals[i] / R_max_abs
        if norm_R >= 0: # Compressione (blu-verde)
            color = plt.cm.viridis(norm_R)
        else: # Trazione (rosso)
            color = plt.cm.plasma(abs(norm_R))
        circle.set_facecolor(color)

        # Etichetta con numero palo e reazione
        ax.text(x_pali[i], y_pali[i], f"{i+1}\n{R_vals[i]:.0f}",
                ha='center', va='center', fontsize=8, color='black',
                bbox=dict(facecolor='white', alpha=0.6, edgecolor='none', pad=1))

    ax.set_title('Geometria in Pianta (Reazioni Statiche Rigide)')
    ax.set_xlabel('x [m]')
    ax.set_ylabel('y [m]')
    ax.set_aspect('equal', adjustable='box')
    ax.grid(True, linestyle='--', alpha=0.6)
    ax.legend(loc='upper right')
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _plot_comparativa_matplotlib(r: Dict[str, Any]) -> bytes:
    """Genera un plot statico comparativo delle reazioni sui pali."""
    st_rig = r['statico']
    se_rig = r['sismico']
    st_fem = r['statico_fem']

    pali_labels = [f"P{i + 1}" for i in range(len(st_rig['x']))]
    x_indices = np.arange(len(pali_labels))
    width = 0.2

    fig, ax = plt.subplots(figsize=(10, 6))

    ax.bar(x_indices - width, st_rig['R'], width, label='Statico (Rigido)', color='steelblue')
    ax.bar(x_indices, st_fem['R'], width, label='Statico (FEM)', color='darkorange')
    ax.bar(x_indices + width, se_rig['R'], width, label='Sismico (Rigido)', color='firebrick')

    qamm = float(r['Qamm_effettiva_palo'])
    ax.axhline(y=qamm, color='red', linestyle='--', label=f'Qamm Eff. = {qamm:.0f} kN')

    ax.set_xlabel('Palo')
    ax.set_ylabel('Reazione [kN]')
    ax.set_title('Confronto Reazioni sui Pali')
    ax.set_xticks(x_indices)
    ax.set_xticklabels(pali_labels)
    ax.legend()
    ax.grid(True, linestyle='--', alpha=0.6)
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


def _plot_stratigrafia_matplotlib(r: Dict[str, Any]) -> bytes:
    """Genera un plot statico della stratigrafia."""
    df = r['stratigrafia']
    fig, ax = plt.subplots(figsize=(6, 8))
    colors = ['#D2B48C', '#DEB887', '#F4A460', '#CD853F', '#A0522D', '#8B4513']

    for i, row in df.iterrows():
        z_t, z_b = row['z_top_m'], row['z_bot_m']
        color = colors[i % len(colors)]
        ax.axhspan(-z_b, -z_t, facecolor=color, edgecolor='black', linewidth=0.8)
        ax.text(0.5, -(z_t + z_b) / 2,
                f"E_ed={row['E_ed_kPa']:.0f} kPa\nφ={row['phi_deg']}° cu={row['cu_kPa']} kPa",
                ha='center', va='center', fontsize=9, color='black')

    ax.set_title("Profilo Stratigrafico")
    ax.set_xlabel("Terreno")
    ax.set_ylabel("Profondità Z [m]")
    ax.set_xlim(0, 1)
    ax.set_ylim(-df['z_bot_m'].max() - 1, 0.5)
    ax.set_xticks([])
    ax.grid(True, linestyle='--', alpha=0.6)
    fig.tight_layout()
    return _fig_to_png_bytes(fig)


# ---------------------------------------------------------------------------
# WORD REPORT
# ---------------------------------------------------------------------------

def create_word_report(d: Any, r: Dict[str, Any]) -> bytes:
    """
    Genera la relazione tecnica in formato Word (.docx) per il plinto su pali.
    Restituisce i bytes del file, senza scrivere su disco.
    """

    def _aggiungi_heading1(doc, testo):
        p = doc.add_paragraph()
        run = p.add_run(testo)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        return p

    def _aggiungi_heading2(doc, testo):
        p = doc.add_paragraph()
        run = p.add_run(testo)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        return p

    def _intestazione_tabella(tabella):
        for cell in tabella.rows[0].cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'E5E7EB')
            cell._tc.get_or_add_tcPr().append(shading)

    def _esito_cell(cell, esito_text: str):
        par = cell.paragraphs[0]
        run = par.add_run(esito_text)
        if esito_text == "VERIFICATO":
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00) # Verde
        elif esito_text == "NON VERIFICATO":
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00) # Rosso
        elif esito_text == "ATTENZIONE":
            run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00) # Arancione
        run.bold = True

    def _aggiungi_immagine(doc, titolo, png_bytes, width_cm=16.0):
        _aggiungi_heading2(doc, titolo)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(png_bytes)
            tmp_path = tmp.name
        try:
            doc.add_picture(tmp_path, width=Cm(width_cm))
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    doc = Document()

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    _aggiungi_heading1(doc, "1. RELAZIONE DI CALCOLO - PLINTO SU PALI")

    _aggiungi_heading2(doc, "1.1 Frontespizio")
    p = doc.add_paragraph()
    p.add_run("Applicazione: ").bold = True
    p.add_run("Plinto su Pali - Analisi Geotecnica e Strutturale")
    p = doc.add_paragraph()
    p.add_run("Normative di riferimento: ").bold = True
    p.add_run("NTC2018 §6.4.3, EC7-1")
    p = doc.add_paragraph()
    p.add_run("Data elaborazione: ").bold = True
    p.add_run(str(date.today()))
    p = doc.add_paragraph()
    p.add_run("Disclaimer: ").bold = True
    p.add_run(
        "I risultati prodotti dall'applicazione devono essere verificati da un ingegnere "
        "abilitato prima di essere utilizzati per qualsiasi scopo progettuale o costruttivo."
    )
    doc.add_paragraph()

    _aggiungi_heading2(doc, "1.2 Sintesi esecutiva")
    sintesi_df = r.get('tabella_sintesi', tabella_sintesi(r))
    if not sintesi_df.empty:
        table = doc.add_table(rows=1, cols=len(sintesi_df.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(sintesi_df.columns):
            table.rows[0].cells[i].text = col_name
        _intestazione_tabella(table)
        for _, row_data in sintesi_df.iterrows():
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)
    
    warnings = r.get("warnings", [])
    if warnings:
        _aggiungi_heading2(doc, "Criticita da esaminare")
        for warning in warnings:
            doc.add_paragraph(str(warning), style="List Bullet")

    _aggiungi_heading2(doc, "1.3 Dati di input")

    _aggiungi_heading2(doc, "1.3.1 Geometria e Materiali")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Parametro"
    table.rows[0].cells[1].text = "Valore"
    _intestazione_tabella(table)
    
    layout_info = r.get('layout_pali', {})
    if layout_info.get('custom'):
        layout_rows = [
            ("Layout pali", "Personalizzato"),
            ("Numero pali", f"{int(layout_info.get('n_totale', len(r['statico']['x'])))}"),
            ("Ingombro gruppo X", f"{float(layout_info.get('larghezza_gruppo_x', 0.0)):.2f} m"),
            ("Ingombro gruppo Y", f"{float(layout_info.get('larghezza_gruppo_y', 0.0)):.2f} m"),
        ]
    else:
        layout_rows = [
            ("Layout pali", "Griglia"),
            ("Numero pali X", f"{d.n_x}"),
            ("Numero pali Y", f"{d.n_y}"),
            ("Interasse pali X", f"{d.interasse_x:.2f} m"),
            ("Interasse pali Y", f"{d.interasse_y:.2f} m"),
        ]

    input_params = [
        ("Larghezza Plinto B", f"{d.B:.2f} m"),
        ("Lunghezza Plinto L", f"{d.L:.2f} m"),
        ("Spessore Plinto", f"{d.spessore_plinto:.2f} m"),
        ("Modulo E calcestruzzo", f"{d.E_cls_MPa:.0f} MPa"),
        *layout_rows,
        ("Diametro Palo", f"{d.diametro_palo:.2f} m"),
        ("Lunghezza Palo", f"{d.lunghezza_palo:.2f} m"),
        ("Falda", f"{d.falda:.2f} m"),
        ("Nq / Nc", f"{d.Nq:.2f} / {d.Nc:.2f}"),
        ("Beta / Alpha", f"{d.beta:.2f} / {d.alpha:.2f}"),
        ("Gamma Sicurezza", f"{d.gamma_sicurezza:.2f}"),
        ("Kh / Kv", f"{d.kh:.2f} / {d.kv:.2f}"),
    ]
    for param, value in input_params:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = value

    _aggiungi_heading2(doc, "1.3.2 Carichi Applicati")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Carico"
    table.rows[0].cells[1].text = "Valore"
    _intestazione_tabella(table)
    
    load_params = [
        ("Carico Verticale N", f"{d.N:.0f} kN"),
        ("Momento Mx", f"{d.Mx:.0f} kNm"),
        ("Momento My", f"{d.My:.0f} kNm"),
    ]
    for param, value in load_params:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = value

    _aggiungi_heading2(doc, "1.3.3 Stratigrafia")
    strat_df = r['stratigrafia']
    if not strat_df.empty:
        table = doc.add_table(rows=1, cols=len(strat_df.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(strat_df.columns):
            table.rows[0].cells[i].text = col_name
        _intestazione_tabella(table)
        for _, row_data in strat_df.iterrows():
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)

    _aggiungi_heading2(doc, "1.3.4 Elaborati grafici di input")
    _aggiungi_immagine(doc, "Figura 1 - Geometria Plinto e Pali (Reazioni Statiche Rigide)", _plot_geometria_matplotlib(d, r))
    _aggiungi_immagine(doc, "Figura 2 - Profilo Stratigrafico", _plot_stratigrafia_matplotlib(r))

    _aggiungi_heading2(doc, "1.4 Risultati di Calcolo")

    _aggiungi_heading2(doc, "1.4.1 Capacità del Palo Singolo e Gruppo")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Parametro"
    table.rows[0].cells[1].text = "Valore"
    _intestazione_tabella(table)
    
    cap_params = [
        ("Qult Palo Isolato", f"{r['Qult_palo']:.0f} kN"),
        ("Qamm Palo Isolato", f"{r['Qamm_palo']:.0f} kN"),
        ("Qamm Effettiva (Gruppo)", f"{r['Qamm_effettiva_palo']:.0f} kN"),
        ("Efficienza Gruppo (eta)", f"{r['efficienza_gruppo']['eta']:.2f}"),
        ("Stato Gruppo", r['efficienza_gruppo']['stato']),
        ("Cedimento Gruppo", f"{r['cedimento_gruppo_mm']:.2f} mm"),
        ("Rigidezza assiale k_v", f"{r['k_v_calcolato']:.0f} kN/m"),
    ]
    for param, value in cap_params:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = value

    _aggiungi_heading2(doc, "1.4.2 Reazioni sui Pali (Confronto Statico/Sismico, Rigido/FEM)")
    comparativa_df = r.get('tabella_pali_comparativa', tabella_pali_comparativa(r))
    if not comparativa_df.empty:
        table = doc.add_table(rows=1, cols=len(comparativa_df.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(comparativa_df.columns):
            table.rows[0].cells[i].text = col_name
        _intestazione_tabella(table)
        for _, row_data in comparativa_df.iterrows():
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)
    
    _aggiungi_immagine(doc, "Figura 3 - Confronto Reazioni sui Pali", _plot_comparativa_matplotlib(r))

    _aggiungi_heading2(doc, "1.4.3 Verifiche di Progetto")
    verifiche_df = r.get('verifiche_df', genera_verifiche_df(d, r))
    if not verifiche_df.empty:
        table = doc.add_table(rows=1, cols=len(verifiche_df.columns))
        table.style = 'Table Grid'
        for i, col_name in enumerate(verifiche_df.columns):
            table.rows[0].cells[i].text = col_name
        _intestazione_tabella(table)
        for _, row_data in verifiche_df.iterrows():
            row = table.add_row().cells
            row[0].text = str(row_data['Verifica'])
            row[1].text = str(row_data['Valore'])
            row[2].text = str(row_data['Limite'])
            row[3].text = str(row_data['D/C'])
            _esito_cell(row[4], str(row_data['Esito']))

    _aggiungi_heading2(doc, "1.4.4 Sollecitazioni per Strut-and-Tie (Plinto Rigido)")
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Sollecitazione"
    table.rows[0].cells[1].text = "Valore"
    _intestazione_tabella(table)
    
    tie_forces = r['strut_and_tie']
    tie_params = [
        ("Trazione X (Tx)", f"{tie_forces['Tx_kN']:.0f} kN"),
        ("Trazione Y (Ty)", f"{tie_forces['Ty_kN']:.0f} kN"),
        ("Altezza utile d", f"{tie_forces['d_utile_m']:.2f} m"),
    ]
    for param, value in tie_params:
        row = table.add_row().cells
        row[0].text = param
        row[1].text = value

    _aggiungi_heading2(doc, "1.5 Note Tecniche")
    notes = r.get("notes", [])
    if notes:
        for note in notes:
            doc.add_paragraph(str(note), style="List Bullet")
    else:
        doc.add_paragraph("Nessuna nota tecnica aggiuntiva.")

    _aggiungi_heading2(doc, "1.6 Disclaimer")
    p_disc = doc.add_paragraph(
        "I risultati prodotti dall'applicazione CivilBox — Plinto su Pali sono elaborati "
        "automaticamente sulla base dei dati inseriti dall'utente. L'applicazione non sostituisce "
        "il giudizio professionale di un ingegnere abilitato. I risultati devono essere verificati "
        "criticamente prima di essere utilizzati per qualsiasi scopo progettuale, costruttivo o "
        "autorizzativo. Anthropic e gli autori dell'applicazione declinano ogni responsabilità "
        "per un uso improprio dei risultati."
    )
    p_disc.runs[0].font.size = Pt(9)

    # Salva in BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
